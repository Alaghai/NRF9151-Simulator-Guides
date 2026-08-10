"""Build the canonical protocol and the Version 8 simulator/server guides."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from micro_protocol import decode_application_packet

ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"
BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
SAMPLES = [
    "AB100032EF0E00010238363133353230363430353037383700010102B513BCFB7CF3D00096010FAC91003B9101AABBCCDDEE01003C03FF01E000",
    "AB100048289A00020238363133353230363430353037383700020202B513BCFB7CF3D0009602B05747FB7D6B16012C020FAC91003B9100112233445502AABBCCDDEE01CCDDEEFFAABB007803FF0096FF",
    "AB100064869400030238363133353230363430353037383700030302B513BCFB7CF3D0009602B05747FB7D6B16012C02B2094EFB7BB90F012C030FAC91003B9100112233445511223344556604AABBCCDDEE01CCDDEEFFAABBAABBCCDDEE03BBCCDDEEFF0000B403FF0096FF",
    "AB100074060A00040238363133353230363430353037383700040402B513BCFB7CF3D0009602B05747FB7D6B16012C02B2094EFB7BB90F012C02B58310FB7C73B000C8040FAC91003B910011223344551122334455660FAC91003B9404AABBCCDDEE01CCDDEEFFAABBAABBCCDDEE03BBCCDDEEFF00012C007803B6FF",
]

PLAIN_ENGLISH_DESCRIPTION = [
    "The Micro tracker uses Bluetooth, GPS/GNSS, LTE-M, and periodic heartbeats together to decide whether the person appears to be in a known safe context and to keep the backend informed. The device has four main reported states inside its heartbeat. 01 means a configured fixed Bluetooth beacon is currently detected. A0 means a configured trusted Bluetooth device is detected. 0A means Bluetooth did not establish the safe context, GPS was used, and the tracker is inside a configured GPS safe zone. 10 means the tracker used GPS and is outside its configured safe zones, so it is in the GPS/LTE tracking state. These values are heartbeat opcodes, while the heartbeat packet itself uses command 0x01.",
    "The local state machine starts with Bluetooth because it can establish the person's context without immediately needing GNSS. At each BLE check, the device first looks for one of its configured fixed beacons. If a recognized beacon is present, the current state becomes Beacon and the heartbeat will report opcode 01. If there is no recognized beacon, it looks for a configured trusted Bluetooth device. If one is present, the state becomes Trusted Device and the heartbeat reports A0. If neither Bluetooth condition succeeds, the device escalates to GNSS. It obtains a position and checks that position against the configured GPS safe zones. Inside a safe zone produces 0A; outside every safe zone produces 10.",
    "The BLE check interval controls how quickly the tracker itself notices that the physical situation has changed. For example, with a 30-second BLE interval, the tracker can determine approximately every 30 seconds whether its home beacon or trusted phone is still nearby. That check is local; it does not inherently require a cellular transmission.",
    "The heartbeat interval controls something different: how frequently the tracker contacts the backend even when the tracker itself believes everything is normal. A heartbeat reports the current state, device health and the complete configured trusted-device list. It also gives the server an opportunity to return a configuration update. If nothing needs changing, the server replies OK. If an update is waiting for that IMEI, the server replies SUP and sends a complete binary command 0x02 configuration packet.",
    "This heartbeat is particularly important when the local tracker cannot know that its definition of “safe” has changed. For example, suppose the person's phone is a trusted device and the person leaves while carrying that phone. The tracker continues detecting the trusted phone and therefore continues to believe the situation is safe. If someone reports the person missing through the backend, the server can queue a new configuration that removes the phone from the trusted-device list. The tracker receives that configuration at its next heartbeat. Therefore, a five-minute heartbeat interval means approximately a five-minute normal maximum wait for the tracker to ask the server for that pending change, assuming cellular connectivity is available. Actual tracking may take somewhat longer because the device still has to process the update, reevaluate its surroundings and obtain GNSS.",
    "Configuration changes use packet command 0x02. This is a full-replacement configuration packet rather than a small individual-setting message. It contains the target IMEI, Update ID, GPS safe zones, fixed beacons, trusted devices and timing settings. A list count of zero clears that list, and the entire packet must validate before any persistent configuration changes.",
    "That gives the backend a way to implement a lost-person workflow without reflashing the device. For example, it could send a new 0x02 configuration with no trusted devices and no GPS safe zones. If the intention is to remove all conditions that can suppress tracking, the fixed beacon list must also be cleared because the normal state machine checks recognized beacons first. After the update is applied, the device should immediately reevaluate its state rather than waiting another complete BLE interval.",
    "Once the tracker determines that it is outside the safe context, the LTE update interval becomes important. This controls how frequently it obtains and reports a fresh position while actively tracking. For example, an LTE interval of 60 seconds could result in approximately one fresh location update every minute while the tracker remains outside. The standalone location packet uses packet command 0x10. This must not be confused with heartbeat opcode 0x10, which simply tells the backend that the heartbeat's current state is GPS + LTE. The canonical V7 protocol defines both command uses separately.",
    "In short, the three timers answer three different questions: BLE interval: “How quickly should the tracker notice its environment changed?” Heartbeat interval: “How frequently should the tracker check in with the backend and give the backend a chance to intervene?” LTE interval: “Once active tracking is required, how frequently should a new location be reported?” Keeping those three responsibilities separate allows fast local detection, reliable remote intervention, and aggressive location tracking when necessary without requiring constant cellular communication during normal safe operation.",
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def base_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, DARK)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    header = section.header.paragraphs[0]
    header.text = "MICRO DATA PACKET | VERSION 8 CANONICAL"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("666666")
    footer = section.footer.paragraphs[0]
    footer.text = "Onomondo Micro Simulator - controlled protocol document"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("666666")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = "Calibri"; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = RGBColor.from_string("0B2545")
    p = doc.add_paragraph(subtitle)
    p.paragraph_format.space_after = Pt(14)
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor.from_string("555555")
    return doc


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        shade(cell, LIGHT)
        set_width(cell, widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True; run.font.size = Pt(9)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = value
            set_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"; run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run.font.size = Pt(8.5)
    shade(paragraph._p.get_or_add_pPr(), "F4F6F9") if False else None


def canonical_doc() -> Document:
    doc = base_doc("Micro Data Packet Canonical Version 8", "State machine, independent timers, heartbeat, LTE-M location, and server-to-device configuration")
    doc.add_heading("1. Status and scope", level=1)
    doc.add_paragraph("This document is the cross-team source of truth. The simulator and TCP-server guides are implementation and operating guides; they must follow this document. It replaces the former command-0x20 typed-TLV settings format. Command 0x02 is positional and full replacement.")
    doc.add_heading("2. Common packet envelope", level=1)
    add_code(doc, "AB | 10 | LENGTH_BE | CRC_BE | SEQUENCE_BE | COMMAND | PAYLOAD")
    add_table(doc, ["Offset", "Field", "Rule"], [
        ["0", "Header", "Fixed 0xAB"], ["1", "Property", "Fixed 0x10"],
        ["2-3", "Length", "uint16 big-endian; counts Command + Payload"],
        ["4-5", "CRC", "CRC-16/XMODEM, initial 0x0000, payload only, big-endian"],
        ["6-7", "Sequence ID", "uint16 big-endian"],
        ["8", "Command", "0x01 heartbeat; 0x02 configuration update; 0x10 LTE-M location"],
        ["9+", "Payload", "Command-specific"],
    ], [0.65, 1.2, 4.75])
    doc.add_paragraph("Length is calculated as total packet bytes minus 8. CRC begins at payload offset 9 and does not include the command. All multi-byte numeric fields are big-endian.")
    doc.add_heading("3. Heartbeat: command 0x01", level=1)
    doc.add_paragraph("The corrected Version 7 heartbeat fields remain, followed by a mandatory complete configured trusted-device registry. This registry is separate from a currently detected beacon or trusted device.")
    add_table(doc, ["Payload field", "Size", "Encoding"], [
        ["IMEI", "15", "ASCII decimal digits"], ["Timestamp", "8", "uint64 UTC Unix milliseconds"],
        ["Battery / isCharging", "1 / 1", "uint8"], ["lastUpdate", "2", "uint16 BE minutes"],
        ["Software / firmware / opcode", "1 / 1 / 1", "uint8"],
        ["locationData", "6 or 12", "6-byte beacon/trusted ID; or 12-byte GPS fields"],
        ["trustedDeviceCount", "1", "0..4"], ["trustedDeviceSlot1..4", "24", "four fixed six-byte slots; unused slots are zero"],
    ], [1.75, 0.65, 4.2])
    doc.add_paragraph("Canonical beacon/trusted heartbeats are 70 bytes with Length 00 3E. Canonical GPS heartbeats are 76 bytes with Length 00 44.")
    doc.add_heading("4. Device state machine and independent timers", level=1)
    doc.add_paragraph("Runtime simulation evaluates local context in this fixed order: configured detected beacon, then configured detected trusted device, then GNSS safe-zone evaluation. A beacon produces opcode 0x01, a trusted device produces 0xA0, GNSS inside a safe zone produces 0x0A, and GNSS outside all safe zones produces 0x10.")
    add_table(doc, ["Timer", "Responsibility", "Packet behavior"], [
        ["bleCheckIntervalSeconds", "Reevaluate local BLE context; escalate to GNSS only if neither configured BLE identity qualifies.", "Does not itself send a heartbeat or LTE location."],
        ["heartbeatIntervalSeconds", "Contact the backend in every state using the latest known state.", "Sends command 0x01; server returns OK newline or SUP newline plus a matching command-0x02 packet."],
        ["LTEupdateIntervalSeconds", "Report fresh location while state is GPS_LTE_OUTSIDE.", "Sends standalone command 0x10 on outside entry and periodically while still outside."],
    ], [1.55, 3.0, 2.0])
    doc.add_paragraph("The simulator starts in configuration-only mode. Runtime simulation is explicitly enabled with `simulation on`; `simulation off` stops all three automatic schedules without changing stored settings. Heartbeat expiry never forces BLE or GNSS reevaluation.")
    doc.add_heading("5. Configuration update: command 0x02", level=1)
    doc.add_paragraph("This is a positional full-replacement update. Validate the entire packet before changing persistent configuration. A zero count clears its corresponding list.")
    add_table(doc, ["Order", "Field", "Encoding"], [
        ["1", "Target IMEI", "15 ASCII decimal digits; must match the receiving device"],
        ["2", "Update ID", "uint16 big-endian; stored after successful application"],
        ["3", "gpsSafeZoneCount", "uint8 0..4"],
        ["4", "safeZone[i]", "int32 latitude e6 + int32 longitude e6 + uint16 radius metres"],
        ["5", "beaconCount / beacon[i]", "uint8 0..4 followed by six-byte fixed beacon identifiers"],
        ["6", "trustedDeviceCount / trustedDevice[i]", "uint8 0..4 followed by six-byte trusted identifiers"],
        ["7", "heartbeatIntervalSeconds", "uint16 BE, 1..65535; one base heartbeat timer"],
        ["8", "LTEupdateIntervalSeconds", "uint16 BE, 1..65535; active only while outside"], ["9", "bleCheckIntervalSeconds", "uint16 BE, 1..65535; same two-byte position formerly named sleepIntervalSeconds"],
        ["10", "SendingUpdate", "one byte: 00 or FF; image transfer is out of scope"],
    ], [0.45, 2.0, 4.15])
    add_code(doc, "Payload = 27 + (10 x safe zones) + (6 x beacons) + (6 x trusted devices)\nLength  = 1 + Payload\nTotal   = 8 + Length")
    doc.add_paragraph("The byte position, width, Length, CRC, and existing regression packet values are unchanged by the BLE-check semantic rename. `sleepIntervalSeconds` is a legacy local CLI alias only; canonical decoding and documentation use `bleCheckIntervalSeconds`.")
    doc.add_heading("6. Application and transport rules", level=1)
    add_bullets(doc, [
        "Configuration is atomic. Bad CRC, Length, IMEI, count, field value, truncated data, invalid flag, or trailing bytes leave the previous persistent configuration unchanged.",
        "The server queues updates by target IMEI. A valid heartbeat receives OK plus newline when no matching update is pending, or SUP plus newline followed by one complete binary command-0x02 packet when one is pending.",
        "TCP is a byte stream. Implementations must buffer split tokens and split packets, retain buffered bytes after a complete packet, and report closure or timeout while a packet is incomplete.",
        "Use binary packet transmission for integration. ASCII-HEX is diagnostic only and must not be confused with binary 0xAB framing.",
    ])
    doc.add_heading("7. Canonical regression vectors and settings", level=1)
    add_table(doc, ["Sample", "Update ID", "Counts (zones/beacons/trusted)", "Heartbeat", "Length / total"], [
        ["1", "1", "1 / 1 / 1", "60 s", "00 32 / 58"], ["2", "2", "2 / 2 / 2", "120 s", "00 48 / 80"],
        ["3", "3", "3 / 3 / 4", "180 s", "00 64 / 108"], ["4", "4", "4 / 4 / 4", "300 s", "00 74 / 124"],
    ], [0.65, 0.8, 2.0, 1.0, 1.55])
    for index, packet in enumerate(SAMPLES, 1):
        decoded = decode_application_packet(bytes.fromhex(packet))
        update = decoded["configuration_update"]
        doc.add_paragraph(f"Sample {index}", style="Heading 3")
        add_code(doc, packet)
        zone_text = "; ".join(
            f"{item['latitude']:.6f}, {item['longitude']:.6f}, {item['radius_m']} m"
            for item in update["safe_zones"]
        ) or "none (clears list)"
        add_table(doc, ["Setting", "Value"], [
            ["Target IMEI / Update ID", f"{update['target_imei']} / {update['update_id']}"],
            ["GPS safe zones", f"{update['gps_safe_zone_count']}: {zone_text}"],
            ["Configured beacons", f"{update['beacon_count']}: {', '.join(update['beacons']) or 'none'}"],
            ["Configured trusted devices", f"{update['trusted_device_count']}: {', '.join(update['trusted_devices']) or 'none'}"],
            ["heartbeatIntervalSeconds", f"{update['heartbeat_interval_seconds']} seconds"],
            ["LTEupdateIntervalSeconds", f"{update['lte_update_interval_seconds']} seconds"],
            ["bleCheckIntervalSeconds", f"{update['ble_check_interval_seconds']} seconds"],
            ["SendingUpdate", f"0x{update['sending_update']:02X}"],
        ], [2.2, 4.35])
    doc.add_heading("8. Plain-English operating description", level=1)
    for paragraph in PLAIN_ENGLISH_DESCRIPTION:
        doc.add_paragraph(paragraph)
    return doc


def protocol_guide() -> Document:
    doc = base_doc("2. Simulator Protocol Decisions and Usage", "Version 8 - configuration mode and explicit runtime simulation")
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph("This implementation guide follows Micro_Data_Packet_V8_State_Machine_and_Settings.docx. It documents simulator behavior, not an alternate wire format.")
    doc.add_heading("2. Decisions implemented", level=1)
    add_table(doc, ["Area", "Implemented decision"], [
        ["Envelope", "AB | 10 | Length BE | CRC BE | Sequence BE | Command | Payload"],
        ["Length", "total packet bytes - 8; Command + Payload only"],
        ["CRC", "CRC-16/XMODEM over payload only, stored big-endian"],
        ["Heartbeat", "Always contains the 25-byte configured trusted-device registry"],
        ["Configuration", "Command 0x02 positional full replacement; target IMEI and Update ID required"],
        ["Operation mode", "Configuration-only by default; `simulation on` explicitly enables automatic behavior"],
        ["State priority", "Configured detected beacon -> configured detected trusted device -> GPS safe zone -> GPS/LTE outside"],
        ["Timers", "BLE check, heartbeat, and LTE location schedules are independent"],
        ["Legacy", "Command 0x20 TLV is rejected and is not emitted"],
    ], [1.4, 4.6])
    doc.add_heading("3. Configuration commands", level=1)
    doc.add_paragraph("config apply <complete_hex_packet> validates and applies the exact packet through the same CRC, Length, IMEI, atomic-validation, and persistence route as a TCP packet. config generate and config set remain convenience tools, but they generate a complete command-0x02 packet from the active full configuration before applying it.")
    add_code(doc, "config apply AB100032EF0E00010238363133353230363430353037383700010102B513BCFB7CF3D00096010FAC91003B9101AABBCCDDEE01003C03FF01E000")
    doc.add_heading("4. Status command", level=1)
    doc.add_paragraph("status prints non-secret runtime and persistent state: IMEI, software and firmware versions, Update ID, lastUpdate, power state, operation mode, current state and heartbeat opcode, currently detected beacon and trusted identity, latest location, every safe zone, every configured beacon, every configured trusted device, all three intervals, SendingUpdate, timer state, transport and TCP state, wire mode, sequence ID, and time simulation state. Configured IDs are explicitly separate from currently detected identities.")
    doc.add_heading("5. Validation and timer behavior", level=1)
    add_bullets(doc, [
        "A valid update replaces all lists and interval values at once; zero count clears a list.",
        "A failed update leaves active persistent settings unchanged.",
        "The successful Update ID is persisted with the configuration.",
        "bleCheckIntervalSeconds is the former sleepIntervalSeconds byte position, with unchanged width, packet Length, and CRC. The old CLI name is a documented legacy alias only.",
        "A configuration change that affects BLE devices, zones, or timer values immediately reevaluates runtime state and replaces the affected schedule without creating a duplicate timer or work item.",
        "Use `simulation on` to begin automatic BLE checks, periodic heartbeats, and LTE locations while outside; use `simulation off` to return to configuration-only mode.",
    ])
    doc.add_heading("6. Regression expectations", level=1)
    add_table(doc, ["Packet", "Expected result"], [
        ["Heartbeat, beacon/trusted opcode", "70 bytes; Length 00 3E; full trusted registry present"],
        ["Heartbeat, GPS opcode", "76 bytes; Length 00 44; full trusted registry present"],
        ["Configuration sample 1", "Update ID 1; 1/1/1 records; heartbeat/LTE/BLE 60/1023/480; flag 00"],
        ["Configuration sample 4", "Update ID 4; 4/4/4 records; heartbeat/LTE/BLE 300/120/950; flag FF; 124 bytes"],
    ], [2.25, 3.75])
    return doc


def server_guide() -> Document:
    doc = base_doc("3. Simulator TCP Server and Command Guide", "Version 8 - canonical command-0x02 update generation and delivery")
    doc.add_heading("1. Server behavior", level=1)
    doc.add_paragraph("The TCP server validates framed simulator packets with the shared canonical decoder. For a valid heartbeat it returns OK newline unless a pending update exists for the heartbeat IMEI. A pending update produces SUP newline followed immediately by one complete binary command-0x02 packet. Invalid input produces ERROR or ERROR:<code> newline. FWUP is recognized only as a deferred token.")
    doc.add_paragraph("The command-0x02 packet is a positional full replacement containing target IMEI, Update ID, GPS safe zones, fixed beacons, trusted devices, heartbeatIntervalSeconds, LTEupdateIntervalSeconds, bleCheckIntervalSeconds, and SendingUpdate. Length BE is uint16 big-endian and counts Command + Payload; CRC is CRC-16/XMODEM over Payload only and is stored big-endian. The bleCheckIntervalSeconds field occupies the former sleepIntervalSeconds byte position; all existing regression HEX remains unchanged.")
    doc.add_heading("2. Stream framing", level=1)
    add_bullets(doc, [
        "Use Length as a big-endian Command + Payload count; total bytes are 8 + Length.",
        "Accept packet and response fragments across arbitrary reads.",
        "Handle SUP and its packet in the same read and retain extra bytes after a packet.",
        "Report a timeout or connection close while an update is incomplete.",
        "The direct LTE and Windows relay return paths use the same response parser.",
    ])
    doc.add_heading("3. Queue a complete configuration", level=1)
    doc.add_paragraph("The update tool requires every field because command 0x02 is a full replacement. List values are concatenated HEX; empty lists use an empty value after the equals sign.")
    add_code(doc, "python micro_update_tool.py queue --imei 861352064050787 --update-id 1 --set heartbeat_interval_seconds=60 --set lte_update_interval_seconds=1023 --set ble_check_interval_seconds=480 --set safe_zones=02B513BCFB7CF3D00096 --set beacon_list=0FAC91003B91 --set trusted_device_list=AABBCCDDEE01 --set sending_update=00")
    doc.add_heading("4. Run and test", level=1)
    add_code(doc, "python -m unittest -v\npython micro_tcp_server.py\npython micro_packet_decoder.py <packet_hex>")
    doc.add_paragraph("The pending-update store is an atomic JSON file keyed by IMEI. A queued record contains its full canonical packet, Update ID, configuration summary, status, and delivery metadata. A packet for one IMEI is never returned in response to a heartbeat from another IMEI.")
    doc.add_heading("5. Security and operating limits", level=1)
    add_bullets(doc, [
        "Do not log or expose SoftSIM profiles, credentials, IRKs, bonding keys, private keys, or passwords.",
        "Six-byte trusted-device and beacon identifiers in this simulator are non-secret test identifiers only.",
        "SendingUpdate indicates 00 or FF only. Firmware-image transfer and installation are intentionally out of scope.",
    ])
    return doc


def write(doc: Document, name: str) -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(DOCS / name)


if __name__ == "__main__":
    # Keep the filenames used by the existing handoff package, while replacing
    # their stale Version 7/TLV contents with the canonical Version 8 guides.
    write(protocol_guide(), "2_Simulator_Protocol_Decisions_and_Usage_v7_SETTINGS_EXTENSION.docx")
    write(server_guide(), "3_Simulator_TCP_Server_and_Command_Guide_v7_SETTINGS_EXTENSION.docx")
